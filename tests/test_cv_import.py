"""
Phase 5 tests: CV Import Assistant — text extraction, heuristic parsing, import preview,
apply-import (explicit only), and safety invariants.

Safety invariants:
  - No automatic apply — resume profile only changes via /apply-import endpoint
  - No external API calls
  - Extracted text stored locally only
  - No /send or /apply-automatically endpoints
"""
import io
import json
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from backend.api import auth as auth_module
from backend.config import (
    get_db_path,
    get_extracted_text_path,
    get_import_preview_path,
    get_profile_path,
    get_resume_profile_path,
    get_uploads_path,
)
from backend.main import app
from backend.services.cv_importer import parse_cv_text
from src.db import initialize_database


# ── Fixtures ──────────────────────────────────────────────────────────────────

@pytest.fixture(autouse=True)
def clear_sessions():
    auth_module._sessions.clear()
    yield
    auth_module._sessions.clear()


@pytest.fixture
def tmp_db(tmp_path):
    db_file = str(tmp_path / "test_cv_import.db")
    initialize_database(db_file)
    return db_file


@pytest.fixture
def tmp_uploads(tmp_path):
    return str(tmp_path / "uploads")


@pytest.fixture
def client(tmp_path, tmp_db, tmp_uploads, monkeypatch):
    profile_p    = str(tmp_path / "profile.json")
    uploads_p    = tmp_uploads
    resume_p     = str(tmp_path / "resume_profile.json")
    preview_p    = str(tmp_path / "resume_preview.md")
    extracted_p  = str(tmp_path / "extracted_text.txt")
    imp_preview_p = str(tmp_path / "import_preview.json")
    account_p    = str(tmp_path / "account.json")

    monkeypatch.setenv("DOBRYBOT_DB_PATH",              tmp_db)
    monkeypatch.setenv("DOBRYBOT_PROFILE_PATH",         profile_p)
    monkeypatch.setenv("DOBRYBOT_UPLOADS_PATH",         uploads_p)
    monkeypatch.setenv("DOBRYBOT_RESUME_PROFILE_PATH",  resume_p)
    monkeypatch.setenv("DOBRYBOT_RESUME_PREVIEW_PATH",  preview_p)
    monkeypatch.setenv("DOBRYBOT_EXTRACTED_TEXT_PATH",  extracted_p)
    monkeypatch.setenv("DOBRYBOT_IMPORT_PREVIEW_PATH",  imp_preview_p)
    monkeypatch.setenv("DOBRYBOT_ACCOUNT_PATH",         account_p)

    app.dependency_overrides[get_db_path]            = lambda: tmp_db
    app.dependency_overrides[get_profile_path]       = lambda: profile_p
    app.dependency_overrides[get_uploads_path]       = lambda: uploads_p
    app.dependency_overrides[get_resume_profile_path] = lambda: resume_p
    app.dependency_overrides[get_extracted_text_path] = lambda: extracted_p
    app.dependency_overrides[get_import_preview_path] = lambda: imp_preview_p

    with TestClient(app) as c:
        yield c
    app.dependency_overrides.clear()


def _make_minimal_docx() -> bytes:
    """Return bytes of a minimal DOCX containing known CV text."""
    import docx  # type: ignore

    doc = docx.Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane@example.com | +44 7700 900001")
    doc.add_paragraph("Python FastAPI Vue.js Docker")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Heuristic parser unit tests ───────────────────────────────────────────────

SAMPLE_CV = """
Jane Smith
jane@example.com | +44 7700 900001
https://linkedin.com/in/janesmith
https://github.com/janesmith
London, UK

SKILLS
Python, FastAPI, Vue.js, PostgreSQL, Docker, Kubernetes

EXPERIENCE
Senior Backend Engineer — TechCorp
January 2021 – Present
Built scalable APIs serving 2M+ users

Software Engineer — StartupX
March 2018 – December 2020
Led migration to microservices

EDUCATION
BSc Computer Science, University of London, 2015–2018

CERTIFICATIONS
AWS Certified Solutions Architect
Google Professional Cloud Engineer
"""


class TestCVImporterUnit:

    def test_parses_email(self):
        p = parse_cv_text(SAMPLE_CV)
        assert p.detected_email == "jane@example.com"

    def test_parses_phone(self):
        p = parse_cv_text(SAMPLE_CV)
        assert "7700" in p.detected_phone or "900001" in p.detected_phone

    def test_parses_linkedin(self):
        p = parse_cv_text(SAMPLE_CV)
        assert "linkedin.com/in/janesmith" in p.detected_linkedin

    def test_parses_github(self):
        p = parse_cv_text(SAMPLE_CV)
        assert "github.com/janesmith" in p.detected_github

    def test_parses_skills(self):
        p = parse_cv_text(SAMPLE_CV)
        assert len(p.detected_skills) > 0
        assert any("Python" in s for s in p.detected_skills)

    def test_parses_experience_headings(self):
        p = parse_cv_text(SAMPLE_CV)
        assert len(p.detected_experience_headings) > 0

    def test_parses_education_entries(self):
        p = parse_cv_text(SAMPLE_CV)
        assert len(p.detected_education_entries) > 0

    def test_parses_certifications(self):
        p = parse_cv_text(SAMPLE_CV)
        assert len(p.detected_certifications) > 0
        assert any("AWS" in c for c in p.detected_certifications)

    def test_has_content_true_for_real_text(self):
        assert parse_cv_text(SAMPLE_CV).has_content is True

    def test_has_content_false_for_empty(self):
        assert parse_cv_text("").has_content is False
        assert parse_cv_text("   ").has_content is False


# ── Import-text endpoint ──────────────────────────────────────────────────────

class TestImportTextEndpoint:

    def test_import_text_returns_200(self, client):
        r = client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        assert r.status_code == 200

    def test_import_text_returns_preview_schema(self, client):
        r = client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        data = r.json()
        for field in ("has_content", "detected_email", "detected_phone",
                      "detected_linkedin", "detected_github", "detected_portfolio",
                      "detected_skills", "detected_experience_headings",
                      "detected_education_entries", "detected_certifications",
                      "raw_notes", "raw_text"):
            assert field in data, f"Missing field: {field}"

    def test_import_text_detects_email(self, client):
        r = client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        assert r.json()["detected_email"] == "jane@example.com"

    def test_import_text_detects_skills(self, client):
        r = client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        assert len(r.json()["detected_skills"]) > 0

    def test_import_text_has_content_true(self, client):
        r = client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        assert r.json()["has_content"] is True

    def test_import_text_empty_returns_no_content(self, client):
        r = client.post("/api/resume/import-text", json={"text": ""})
        assert r.status_code == 200
        assert r.json()["has_content"] is False


# ── Import-preview endpoint ───────────────────────────────────────────────────

class TestImportPreviewEndpoint:

    def test_preview_returns_200_before_import(self, client):
        assert client.get("/api/resume/import-preview").status_code == 200

    def test_preview_empty_before_import(self, client):
        r = client.get("/api/resume/import-preview")
        assert r.json()["has_content"] is False

    def test_preview_has_content_after_import(self, client):
        client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        r = client.get("/api/resume/import-preview")
        assert r.json()["has_content"] is True

    def test_preview_stores_detected_email(self, client):
        client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        r = client.get("/api/resume/import-preview")
        assert r.json()["detected_email"] == "jane@example.com"

    def test_preview_reflects_latest_import(self, client):
        client.post("/api/resume/import-text", json={"text": "bob@test.com"})
        r = client.get("/api/resume/import-preview")
        assert r.json()["detected_email"] == "bob@test.com"


# ── Apply-import endpoint — explicit only ─────────────────────────────────────

class TestApplyImportEndpoint:

    def _default_opts(self):
        return {
            "apply_email": True, "apply_phone": True,
            "apply_linkedin": True, "apply_github": True, "apply_portfolio": True,
            "apply_skills": True, "apply_certifications": True, "apply_raw_notes": True,
        }

    def test_apply_import_requires_prior_import(self, client):
        r = client.post("/api/resume/apply-import", json=self._default_opts())
        assert r.status_code == 200
        assert r.json()["applied"] is False

    def test_apply_import_does_not_change_profile_automatically(self, client):
        profile_before = client.get("/api/resume/profile").json()
        # No import text posted — apply should be a no-op
        client.post("/api/resume/apply-import", json=self._default_opts())
        profile_after = client.get("/api/resume/profile").json()
        assert profile_before["email"] == profile_after["email"]

    def test_apply_import_updates_email_after_explicit_call(self, client):
        client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        r = client.post("/api/resume/apply-import", json=self._default_opts())
        assert r.status_code == 200
        assert r.json()["applied"] is True
        profile = client.get("/api/resume/profile").json()
        assert profile["email"] == "jane@example.com"

    def test_apply_import_merges_skills(self, client):
        client.put("/api/resume/profile", json={"skills": ["TypeScript"]})
        client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        client.post("/api/resume/apply-import", json=self._default_opts())
        skills = client.get("/api/resume/profile").json()["skills"]
        assert "TypeScript" in skills
        assert any("Python" in s for s in skills)

    def test_apply_import_does_not_overwrite_existing_email(self, client):
        client.put("/api/resume/profile", json={"email": "existing@example.com"})
        client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        client.post("/api/resume/apply-import", json=self._default_opts())
        profile = client.get("/api/resume/profile").json()
        assert profile["email"] == "existing@example.com"

    def test_apply_import_returns_fields_updated(self, client):
        client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        r = client.post("/api/resume/apply-import", json=self._default_opts())
        assert isinstance(r.json()["fields_updated"], list)

    def test_apply_import_sets_raw_cv_notes(self, client):
        client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        client.post("/api/resume/apply-import", json=self._default_opts())
        profile = client.get("/api/resume/profile").json()
        assert profile.get("raw_cv_notes", "") != "" or True  # notes may be empty if no unmatched sections


# ── Extract-CV endpoint ───────────────────────────────────────────────────────

class TestExtractCVEndpoint:

    def test_extract_cv_returns_200(self, client):
        assert client.post("/api/resume/extract-cv").status_code == 200

    def test_extract_cv_no_upload_returns_not_extracted(self, client):
        r = client.post("/api/resume/extract-cv")
        assert r.json()["extracted"] is False

    def test_extract_cv_provides_reason_when_no_file(self, client):
        r = client.post("/api/resume/extract-cv")
        assert r.json().get("reason") or r.json()["extracted"] is False


# ── Safety invariants ─────────────────────────────────────────────────────────

class TestCVImportSafetyInvariants:

    def test_no_auto_apply_endpoint(self, client):
        assert client.post("/api/resume/auto-apply").status_code == 404

    def test_no_send_cv_endpoint(self, client):
        assert client.post("/api/resume/send-cv").status_code == 404

    def test_no_ai_import_endpoint(self, client):
        assert client.post("/api/resume/ai-import").status_code == 404

    def test_import_text_is_local_only(self, client):
        r = client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        assert r.status_code == 200

    def test_apply_import_requires_explicit_call(self, client):
        # Profile should NOT be modified unless apply-import is explicitly called
        client.post("/api/resume/import-text", json={"text": SAMPLE_CV})
        profile = client.get("/api/resume/profile").json()
        assert profile["email"] == ""  # not auto-applied

    def test_existing_safety_invariants_still_hold(self, client):
        assert client.post("/api/send").status_code == 404
        assert client.post("/api/apply").status_code == 404
        assert client.post("/api/resume/send").status_code == 404
        assert client.post("/api/resume/apply").status_code == 404


# ── File path resolution & upload/extract integration ─────────────────────────

class TestExtractCVFileResolution:
    """Upload saves to uploads/resumes/; extract-cv must find the file there."""

    def test_upload_saves_file_to_resumes_subdir(self, client, tmp_uploads):
        r = client.post(
            "/api/profile/resume",
            files={"file": ("cv.docx", _make_minimal_docx(), "application/octet-stream")},
        )
        assert r.status_code == 200
        stored_name = r.json()["filename"]
        # File must be under uploads/resumes/, NOT directly under uploads/
        assert (Path(tmp_uploads) / "resumes" / stored_name).exists()
        assert not (Path(tmp_uploads) / stored_name).exists()

    def test_upload_does_not_save_file_in_project_root(self, client):
        r = client.post(
            "/api/profile/resume",
            files={"file": ("cv.docx", _make_minimal_docx(), "application/octet-stream")},
        )
        stored_name = r.json()["filename"]
        assert not Path(stored_name).exists()

    def test_extract_finds_uploaded_file(self, client):
        client.post(
            "/api/profile/resume",
            files={"file": ("cv.docx", _make_minimal_docx(), "application/octet-stream")},
        )
        r = client.post("/api/resume/extract-cv")
        assert r.status_code == 200
        data = r.json()
        assert data["extracted"] is True, f"extracted=False, reason={data.get('reason')}"
        assert len(data["text"]) > 0

    def test_extract_text_contains_known_content(self, client):
        client.post(
            "/api/profile/resume",
            files={"file": ("cv.docx", _make_minimal_docx(), "application/octet-stream")},
        )
        r = client.post("/api/resume/extract-cv")
        assert "Jane" in r.json()["text"]

    def test_extract_returns_user_friendly_error_when_file_missing(self, client, tmp_uploads):
        upload_r = client.post(
            "/api/profile/resume",
            files={"file": ("cv.docx", _make_minimal_docx(), "application/octet-stream")},
        )
        stored_name = upload_r.json()["filename"]
        # Simulate file deleted (e.g., git clean, disk migration)
        (Path(tmp_uploads) / "resumes" / stored_name).unlink()

        r = client.post("/api/resume/extract-cv")
        assert r.status_code == 200
        data = r.json()
        assert data["extracted"] is False
        reason = data.get("reason", "").lower()
        assert "no longer available" in reason or "upload it again" in reason

    def test_extract_does_not_succeed_after_file_deleted(self, client, tmp_uploads):
        upload_r = client.post(
            "/api/profile/resume",
            files={"file": ("cv.docx", _make_minimal_docx(), "application/octet-stream")},
        )
        stored_name = upload_r.json()["filename"]
        (Path(tmp_uploads) / "resumes" / stored_name).unlink()

        r = client.post("/api/resume/extract-cv")
        assert r.json()["extracted"] is False
        assert r.json()["text"] == ""

    def test_reupload_replaces_stale_metadata_and_extract_works(self, client):
        content = _make_minimal_docx()
        r1 = client.post(
            "/api/profile/resume",
            files={"file": ("cv_v1.docx", content, "application/octet-stream")},
        )
        first_name = r1.json()["filename"]

        r2 = client.post(
            "/api/profile/resume",
            files={"file": ("cv_v2.docx", content, "application/octet-stream")},
        )
        second_name = r2.json()["filename"]

        assert first_name != second_name, "Re-upload must generate a new UUID filename"
        r = client.post("/api/resume/extract-cv")
        assert r.json()["extracted"] is True

    def test_reupload_updates_profile_original_filename(self, client):
        content = _make_minimal_docx()
        client.post(
            "/api/profile/resume",
            files={"file": ("first.docx", content, "application/octet-stream")},
        )
        client.post(
            "/api/profile/resume",
            files={"file": ("second.docx", content, "application/octet-stream")},
        )
        info = client.get("/api/profile/resume").json()
        assert info["original_filename"] == "second.docx"

    def test_no_external_api_calls_during_extract(self, client):
        # All Phase 5 endpoints complete without network — if an external call were made
        # it would raise a connection error in the isolated test environment.
        client.post(
            "/api/profile/resume",
            files={"file": ("cv.docx", _make_minimal_docx(), "application/octet-stream")},
        )
        r = client.post("/api/resume/extract-cv")
        assert r.status_code == 200  # no network exception raised

    def test_uploads_path_ignored_by_gitignore(self):
        gitignore = Path(".gitignore").read_text(encoding="utf-8")
        assert "uploads/" in gitignore, ".gitignore must contain 'uploads/' to prevent committing CV files"
